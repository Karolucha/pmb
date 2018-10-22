from datetime import datetime, timedelta

from django.http import HttpResponse
from docx import Document

from adminek.views.generic_views import BaseGenericView
from django.shortcuts import render, redirect
from app.models import Intentions, IntentionWeek, MassSchema


class IntentionView(BaseGenericView):
    messes = {}

    def __init__(self, **kwargs):
        self.days = ['Niedziela', 'Poniedziałek', 'Wtorek', 'Środa', 'Czwartek', 'Piątek', 'Sobota']
        super().__init__(**kwargs)

    def get_for_single(self, model_class, id_object):
        object_context = IntentionWeek.objects.filter(id=id_object).prefetch_related('intentions_set')[0]
        intentions = []
        for intention in object_context.intentions_set.all().order_by('date', 'hour'):
            intentions.append(intention)
        return {
            'intentionweek': object_context,
            'intentions': intentions
        }

    def get(self, request, *args, **kwargs):
        method = kwargs.get('method', 'get')
        if method == 'edit':
            self.prepare_ready_intentions(kwargs['pk'])
            return render(request, self.template_name, self.context)
        elif method == 'download':
            self.prepare_ready_intentions(kwargs['pk'])
            document = IntentionWord(self.context['intention_week_start'], self.context['intentions'])
            return document.download_docx()
        else:
            return super().get(request, *args, **kwargs)

    def prepare_ready_intentions(self, intention_id):
        self.context = {}
        intentions = []
        intention_week = self.get_for_single(self.model_class, intention_id)
        date = intention_week['intentionweek'].week
        self.context['intention_week_start'] = intention_week['intentionweek'].week
        self.context['intention_id'] = intention_week['intentionweek'].id
        i = 0
        intentions_day = []
        for intention in intention_week['intentions']:
            if intention.date != date:
                intentions.append({
                    'day': self.days[i],
                    'intentions': intentions_day.copy()})
                i += 1
                date = intention.date
                intentions_day = []
            intentions_day.append(intention)

        intentions.append({
            'day': self.days[i],
            'intentions': intentions_day.copy()})
        self.template_name = 'others/intentionweek_edit.html'
        self.context['intentions'] = intentions

    def post(self, request, *args, **kwargs):
        self.context = {}
        params = dict(request.POST)
        if 'stepDay' in params.keys():
            self.context['object'] = {}
            # self.context['object']['date'] = request.POST['date']
            old_week = IntentionWeek.objects.filter(week=request.POST['date'])
            if old_week.exists():
                self.context['intention_week_start'] = old_week[0].week
                return redirect('intentionweek', method='edit', pk=old_week[0].id, object_name='intentionweek')
            self.context['intention_week_start'] = request.POST['date']
            intentions = [self.transform_mass_schema_to_intentions(True, self.days[0], 0)]
            for index, day in enumerate(self.days[1:]):
                intentions.append(self.transform_mass_schema_to_intentions(False, day, index + 1))
            self.context['intentions'] = intentions
            # print('context ', self.context, self.template_name)
            return render(request, 'others/intentions_form.html', self.context)
        else:
            return super().post(request, *args, **kwargs)

    def transform_mass_schema_to_intentions(self, is_sunday, day, index):
        now = datetime.now().date()
        intentions = []
        mass_schemas = MassSchema.objects.filter(
            season_start__lt=now, season_end__gt=now, sunday=is_sunday).prefetch_related('hour_set')
        hours_list_in_day = [hour.hour for schema in mass_schemas for hour in schema.hour_set.all()]
        hours_list_in_day.sort()
        date = (datetime.strptime(self.context['intention_week_start'], '%Y-%m-%d') + timedelta(days=index)).date()
        for hour in hours_list_in_day:
            intention = Intentions(hour=hour, date=date)
            intentions.append(intention)
        return {
            'day': day,
            'intentions': intentions}

    def create(self, request, *args, **kwargs):
        params = dict(request.POST)
        intention_week = IntentionWeek(week=request.POST['date'])
        intention_week.save()

        for day in self.days:
            hours_for_day = {key: value for key, value in params.items() if key.startswith(day+'-hour')}
            for hour_key in hours_for_day:
                hour = hour_key.split('-')[-1]
                intention = Intentions(date=params[day+'-date-'+hour][0], hour=hour,
                                       title=params[day+'-content-'+hour][0], week=intention_week)
                intention.save()

        return redirect('detail', method='list', object_name='intentionweek')

    def edit(self, request, *args, **kwargs):
        params = dict(request.POST)
        for day in self.days:
            hours_for_day = {key: value[0] for key, value in params.items() if key.startswith(day+'-hour')}
            for hour_key in hours_for_day:
                hour = hour_key.split('-')[-1]
                intention = Intentions.objects.filter(id=params[day+'-id-'+hour][0])[0]
                intention.title = params[day+'-content-'+hour][0]
                intention.save()


class IntentionWord:
    def __init__(self, date, announcements):
        self.day = str(date)
        self.announcements = announcements

    def download_docx(self):
        document = Document()
        document.add_heading('Intencje na tydzień ' + self.day, 0)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dzień'
        hdr_cells[1].text = 'Godzina'
        hdr_cells[2].text = 'Intencja'

        for intention_day in self.announcements:
            for intention in intention_day['intentions']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(intention_day['day'])
                row_cells[1].text = str(intention.hour)[:-3]
                row_cells[2].text = intention.title
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=intencje.docx'
        document.save(response)

        return response
