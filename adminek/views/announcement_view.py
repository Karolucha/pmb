import os
from adminek.views.generic_views import BaseGenericView
from django.shortcuts import render, redirect
from django.core.files import File
from django.http import HttpResponse
from app.models import WeekAnnouncement, Announcement
from docx import Document


class AnnouncementView(BaseGenericView):
    messes = {}

    def get_for_single(self, model_class, id_object):
        object_context = model_class.objects.filter(id=id_object).prefetch_related('announcement_set')[0]
        hours_set = []
        for msg in object_context.announcement_set.all().order_by('id'):
            hours_set.append(msg)
        return {
            'schema': object_context,
            'announcements': hours_set
        }

    def get(self, request, *args, **kwargs):
        method = kwargs.get('method', 'get')
        if method == 'download':
            object_context = WeekAnnouncement.objects.filter(id=kwargs['pk']).prefetch_related('announcement_set')[0]
            document = DocumentWord(object_context.date, object_context.announcement_set.all().order_by('id'))
            return document.download_docx()
        else:
            return super().get(request, *args, **kwargs)

    def create(self, request, *args, **kwargs):
        schema = WeekAnnouncement(date=request.POST['date'])
        schema.save()
        self.add_new_announcements(request, schema)

        return redirect('detail', method='list', object_name='weekannouncement')

    def edit(self, request, *args, **kwargs):
        anothers = dict(request.POST)
        schema = WeekAnnouncement.objects.get(id=kwargs['pk'])
        schema.save()
        self.add_new_announcements(request, schema)
        self.messes = {}

        for name, value in anothers.items():
            if name.startswith('old-'):
                row_number = name.split('-')[-1]
                announcement = Announcement.objects.get(id=row_number)
                announcement.content = value[0]
                announcement.save()
            if name == 'deletions':
                for hour_element_id in value:
                    hour_id = hour_element_id.split('-')[-1]
                    hour = Announcement.objects.get(id=hour_id)
                    hour.delete()
        schema.save()

    def add_new_announcements(self, request, schema):
        anothers = dict(request.POST)
        for name, value in anothers.items():
            if name.startswith('n-a-'):
                announcement = Announcement(content=value[0],
                                            week_announcment=schema)
                announcement.save()


class DocumentWord:
    def __init__(self, date, announcements):
        self.day = str(date)
        self.announcements = announcements

    def create(self):
        document = Document()
        document.add_heading('Ogłoszenia duszpasterskie ' + self.day, 0)
        for announcement in self.announcements:
            document.add_paragraph(
                announcement.content, style='ListNumber'
            )
        document.add_page_break()

        document.save('demo.docx')

    def download_docx(self):
        document = Document()
        document.add_heading('Ogłoszenia duszpasterskie ' + self.day, 0)
        for announcement in self.announcements:
            document.add_paragraph(
                announcement.content, style='ListNumber'
            )
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=ogloszenia.docx'
        document.save(response)

        return response
